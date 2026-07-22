"""
parser_core.py

Deterministic ELA lesson .docx parser (adapted from AI INSTRUCTIONS/ela_to_row.py).
Sections are detected by BOLD label text against a whitelist -- extraction is
verbatim by construction: text is sliced, never regenerated.

Unknown bold labels are collected; a judge (Ollama/Gemini) may map them onto
canonical columns. Nothing is silently dropped: unmatched sections keep their
own column, loose text goes to _misc.

Author: AbhishekAEDan
"""
__author__ = "AbhishekAEDan"

import os
import re
import zipfile

from docx import Document

# normalized bold-label prefix -> output column name
LABELS = [
    ("my goal",                    "My Goal"),
    ("by the end of this lesson",  "Lesson Objectives"),
    ("bridge back",                "Bridge Back"),
    ("vocab vault",                "Vocab Vault"),
    ("vocabulary vault",           "Vocab Vault"),
    ("mistake spotter",            "Mistake Spotter"),
    ("main comprehension text",    "Comprehension Title"),
    ("before reading",             "Before Reading"),
    ("after reading",              "After Reading"),
    ("comprehension questions",    "Comprehension Questions"),
    ("main text",                  "Main Text"),
    ("main lesson",                "Main Lesson"),
    ("mini practice",              "Mini Practice"),
    ("check for understanding",    "Check for Understanding"),
    ("sims interactive",           "Interactive / Simulation"),
    ("development notes",          "Interactive / Simulation"),
    ("portfolio project",          "Portfolio Project"),
    ("end-of-unit inventory",      "End-of-Unit Inventory"),
    ("my learning checklist",      "My Learning Checklist"),
    ("quick recap",                "Quick Recap"),
    ("reflect",                    "Reflect"),
]

CANONICAL_COLUMNS = sorted({col for _, col in LABELS})

IMAGE_PATTERNS = [
    re.compile(r'\[IMAGE:\s*(.+?)\]', re.I | re.S),
    re.compile(r'Image Suggestion:\s*(.+)', re.I),
    re.compile(r'Simulation Idea:\s*(.+)', re.I),
]

ORDER = ["_file", "_doc_type", "Unit Number", "Lesson Title", "My Goal",
         "Lesson Objectives", "Bridge Back", "Vocab Vault", "Mistake Spotter",
         "Comprehension Title", "Before Reading", "Main Text",
         "Comprehension Questions", "Comprehension Questions Answer Key",
         "After Reading", "Main Lesson",
         "Mini Practice Title", "Mini Practice Intro", "Practice A", "Practice B",
         "Practice C", "What to look for", "Mini Practice Answer Key",
         "Check for Understanding", "Check for Understanding Answer Key",
         "Interactive / Simulation", "Portfolio Project", "End-of-Unit Inventory",
         "My Learning Checklist", "Quick Recap", "Reflect",
         # unit-level (Intro / End-of-Unit / Portfolio) columns
         "Unit Title", "Top Content", "At End Of Unit Checklist",
         "Unit Learning Checklist", "Unit Quick Recap", "Unit Reflect",
         "Project Title", "Purpose", "Portfolio Link", "What You Will Create",
         "Materials", "Steps", "Template How To Fill", "Example",
         "Common Mistakes And Fixes", "Submission Checklist", "Optional Extension",
         "Motivation Reflection",
         "has_image", "image_missing", "image_note", "_misc"]


def bold_fraction(p):
    tot = bold = 0
    for r in p.runs:
        t = r.text or ""
        if not t.strip():
            continue
        tot += len(t)
        if r.bold:
            bold += len(t)
    return (bold / tot) if tot else 0.0


def leading_bold(p):
    out = []
    for r in p.runs:
        if r.text and r.text.strip():
            if r.bold:
                out.append(r.text)
            else:
                break
    return "".join(out).strip()


def after_leading_bold(p):
    started = False
    out = []
    for r in p.runs:
        t = r.text or ""
        if not started:
            if t.strip() and r.bold:
                continue
            started = True
        out.append(t)
    return "".join(out).strip()


def is_list_item(p):
    pPr = p._p.pPr
    return pPr is not None and pPr.numPr is not None


# bold lead-ins that are content structure, never section headers
CONTENT_PREFIXES = re.compile(
    r'^(step\s*\d|[a-c]\)|example\b|question\b|answer:|note\b|common mistakes\b|'
    r'numbered steps\b|submission checklist\b|vocab preview\b|\[|\d)', re.I)

BULLET_CHARS = "•·●▪�-–—o \t"


def clean_label(s):
    return s.lstrip(BULLET_CHARS).strip()


# leading list glyphs Word bakes into runs (wingdings bullet, dots, etc.)
_BULLET_RE = re.compile('^[•·●▪⁃]+[ 	]*')


def strip_bullet(line):
    return _BULLET_RE.sub('', line)


def norm_label(s):
    s = clean_label(s)
    s = re.split(r'[—:(]', s)[0]
    return s.rstrip(" .!…-").lower().strip()


# fallback: label families matched anywhere in the text
CONTAINS_LABELS = [
    ("glossary",           "Glossary"),
    ("recap",              "Quick Recap"),
    ("reflect",            "Reflect"),
    ("learning checklist", "My Learning Checklist"),
    ("i will be able to",  "Lesson Objectives"),
]


def match_label(text):
    n = norm_label(text)
    for key, col in LABELS:
        if n.startswith(key):
            return col
    full = clean_label(text).lower()
    for key, col in CONTAINS_LABELS:
        if key in full:
            return col
    return None


# a comprehension passage is titled with a genre lead, e.g.
#   Prose Text: "..."   Poem: "..."   Drama Scene: "..."
_COMP_TITLE_RE = re.compile(
    r'^(prose text|prose|poem|poetry(?:\s+text)?|drama scene|drama|play|'
    r'non-?fiction(?:\s+text)?|informational text|expository text|'
    r'main comprehension text)\b\s*:', re.I)


def find_comprehension_title(doc):
    """Return the genre-titled passage heading (e.g. 'Poem: \"...\"') if the
    document contains one -- used when the title line was not captured as a
    section (it is often number-prefixed or non-bold)."""
    for p in doc.paragraphs:
        for ln in p.text.split("\n"):
            s = strip_bullet(ln.strip())
            m = _COMP_TITLE_RE.match(s)
            if not m:
                continue
            # "Main Comprehension Text:" is a label, the title is what follows
            if m.group(1).lower() == "main comprehension text":
                rest = s[m.end():].strip()
                if rest:
                    return rest
                continue
            return s
    return ""


def subtitle_after_dash(text):
    m = re.split(r'\s*—\s*', text, maxsplit=1)
    return m[1].strip() if len(m) > 1 else ""


def split_sections(doc, unknown_labels):
    """Return (title, [(column_or_rawlabel, [lines])]).
    Unknown bold labels are recorded in unknown_labels (set) and used
    verbatim as the column name pending judge remap."""
    paras = [p for p in doc.paragraphs if p.text.strip()]
    title = paras[0].text.strip() if paras else ""
    sections = []
    current = None
    last_real_col = None
    in_mini = False          # inside the Mini Practice section
    in_main = False          # inside the Main Lesson section
    for p in paras[1:]:
        raw = p.text.strip()
        txt = strip_bullet(raw)
        bf = bold_fraction(p)
        if bf >= 0.9:
            label_txt = txt
        elif not is_list_item(p):
            label_txt = leading_bold(p)
        else:
            label_txt = ""
        # bullet-style or structural bold lead-ins are content, not headers
        if label_txt:
            cl = clean_label(label_txt)
            is_bulleted = bool(re.match(r'^[•·●▪�]|^[o\-–—]\s', txt.lstrip()))
            if not cl or CONTENT_PREFIXES.match(cl) or is_bulleted:
                label_txt = ""
            else:
                label_txt = cl
        col = match_label(label_txt) if label_txt else None
        n = norm_label(label_txt) if label_txt else ""
        # "Answer Key:", "Answers:", but also suffixed forms like
        # "Check for Understanding - Answer Key:"
        is_answer_key = (n.startswith("answer key") or n == "answers"
                         or n.endswith("answer key"))

        # unknown bold lead ending with ':' is a sub-label inside the current
        # section (e.g. "Steps:", "What to look for:", "Correct state:"),
        # not a new section
        if (col is None and not is_answer_key and label_txt
                and label_txt.rstrip().endswith(":")):
            label_txt = ""
            col = None

        if is_answer_key:
            akcol = (last_real_col + " Answer Key") if last_real_col else "Answer Key"
            # the answer text is sometimes on following paragraphs, sometimes
            # inline after the bold "Answer Key:" label in the same paragraph
            rest = after_leading_bold(p) if bf < 0.9 else ""
            current = (akcol, [rest] if rest.strip() else [])
            sections.append(current)
            continue

        # while inside Mini Practice, bold sub-titles (e.g. a completed-example
        # heading like "DeAndre's Hidden Insecurity") are content, not new
        # sections -- keeps a), b), c) and the answer key in one block
        if col is None and label_txt and bf >= 0.9 and in_mini:
            if current is not None:
                current[1].append(txt)
            continue

        # likewise inside Main Lesson: bold sub-headings ("Characters in
        # Drama", "Step 2: ...") are lesson content, not new sections --
        # they stay in the Main Lesson text so the Main Lesson sheet can
        # split them into Top Content / Step columns
        if col is None and label_txt and bf >= 0.9 and in_main:
            if current is not None:
                current[1].append(txt)
            continue

        # a full-line bold label that isn't whitelisted: treat as unknown section
        if col is None and label_txt and bf >= 0.9 and len(label_txt) <= 60:
            unknown_labels.add(label_txt)
            current = ("?" + label_txt, [])
            sections.append(current)
            last_real_col = "?" + label_txt
            sub = subtitle_after_dash(txt)
            if sub:
                current[1].append(f"[title] {sub}")
            continue

        if col:
            current = (col, [])
            sections.append(current)
            last_real_col = col
            in_mini = (col == "Mini Practice")
            in_main = (col == "Main Lesson")
            if bf < 0.9:
                rest = after_leading_bold(p)
                if rest:
                    current[1].append(rest)
            else:
                sub = subtitle_after_dash(txt)
                if sub:
                    current[1].append(f"[title] {sub}")
            continue

        if current is not None:
            current[1].append(txt)
        else:
            sections.append(("_misc", [txt]))
            current = sections[-1]
    return title, sections


def fan_mini_practice(lines):
    out = {"Mini Practice Title": "", "Mini Practice Intro": "",
           "Practice A": "", "Practice B": "", "Practice C": "",
           "What to look for": ""}
    buckets = {"intro": [], "a": [], "b": [], "c": [], "wtlf": []}
    cur = "intro"
    for ln in lines:
        low = ln.lower()
        if ln.startswith("[title] "):
            out["Mini Practice Title"] = ln[len("[title] "):]
            continue
        # the marker sometimes trails a short lead-in ("End with: What to
        # look for: ...") -- allow a small prefix, but not a mid-sentence hit
        wtlf = re.match(r'^.{0,12}what to look for', low)
        if wtlf:
            cur = "wtlf"
            buckets["wtlf"].append(ln[low.find("what to look for"):])
            continue
        m = re.match(r'^([a-cA-C])\)\s', ln)
        if m:
            cur = m.group(1).lower()
            buckets[cur].append(ln)
            continue
        buckets[cur].append(ln)
    out["Mini Practice Intro"] = "\n".join(buckets["intro"]).strip()
    out["Practice A"] = "\n".join(buckets["a"]).strip()
    out["Practice B"] = "\n".join(buckets["b"]).strip()
    out["Practice C"] = "\n".join(buckets["c"]).strip()
    out["What to look for"] = "\n".join(buckets["wtlf"]).strip()
    return out


def detect_images(path, all_text):
    notes = []
    embedded = False
    try:
        with zipfile.ZipFile(path) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            embedded = len(media) > 0
            if "word/document.xml" in z.namelist():
                xml = z.read("word/document.xml").decode("utf-8", "ignore")
                for descr in re.findall(r'descr="([^"]+)"', xml):
                    if descr.strip():
                        notes.append(descr.strip())
    except Exception:
        pass
    for pat in IMAGE_PATTERNS:
        for m in pat.finditer(all_text):
            notes.append(m.group(1).strip()[:160])
    has = embedded or bool(notes)
    return {
        "has_image": "yes" if has else "no",
        "image_missing": "yes" if has else "no",
        "image_note": " || ".join(dict.fromkeys(notes)),
    }


# ---------- document-type routing (author: AbhishekAEDan) ----------
#
# The generic bold-label parser above handles ordinary sub-unit lessons.
# Three other document kinds have their own fixed internal layout and are
# parsed by dedicated functions so their fields land in the right sheet:
#   Unit N Introduction      -> Unit Intro Page
#   End of Unit N            -> End Of Unit Portfolio Doc
#   End of Unit N Portfolio  -> Overall Portfolio Project

_RE_UNIT_INTRO = re.compile(r'^unit\s+(\d+)\s+introduction', re.I)
_RE_EOU_PORTFOLIO = re.compile(r'^end of unit\s+(\d+)\s+portfolio', re.I)
_RE_EOU = re.compile(r'^end of unit\s+(\d+)', re.I)


def classify_doc(filename):
    base = os.path.splitext(os.path.basename(filename))[0]
    if _RE_UNIT_INTRO.match(base):
        return "unit_intro"
    if _RE_EOU_PORTFOLIO.match(base):
        return "end_of_unit_portfolio"
    if _RE_EOU.match(base):
        return "end_of_unit"
    return "lesson"


def _unit_from_name(filename, patterns):
    base = os.path.basename(filename)
    for pat in patterns:
        m = pat.match(os.path.splitext(base)[0])
        if m:
            return m.group(1)
    return ""


def _nonempty_paras(doc):
    return [p for p in doc.paragraphs if p.text.strip()]


def parse_unit_intro(doc, path):
    """Unit N Introduction -> Unit Intro Page fields.
    Title line 'UNIT N — Title'; everything up to the bold
    'At the End of Unit N, I Will Be Able To' line is Top Content; that
    heading and its checklist form the end-of-unit checklist column."""
    paras = _nonempty_paras(doc)
    row = {"_file": os.path.basename(path)}
    row["Unit Number"] = _unit_from_name(path, [_RE_UNIT_INTRO])
    title_line = paras[0].text.strip() if paras else ""
    m = re.match(r'^\s*unit\s+\d+\s*[—–-]\s*(.*)$', title_line, re.I)
    row["Unit Title"] = m.group(1).strip() if m else title_line

    top, checklist = [], []
    in_checklist = False
    for p in paras[1:]:
        txt = strip_bullet(p.text.strip())
        if not in_checklist and re.match(r'^at the end of unit\b', txt, re.I):
            in_checklist = True
        (checklist if in_checklist else top).append(txt)
    row["Top Content"] = "\n".join(top).strip()
    row["At End Of Unit Checklist"] = "\n".join(checklist).strip()
    return row


def parse_end_of_unit(doc, path):
    """End of Unit N -> checklist / quick recap / reflect columns.
    Split on the three bold section heads."""
    paras = _nonempty_paras(doc)
    row = {"_file": os.path.basename(path)}
    row["Unit Number"] = _unit_from_name(path, [_RE_EOU])

    buckets = {"checklist": [], "recap": [], "reflect": []}
    cur = "checklist"
    for p in paras:
        txt = strip_bullet(p.text.strip())
        low = txt.lower()
        # headings vary in word order across units and may share a paragraph
        # with their body, so test the first line only
        first = low.split("\n", 1)[0]
        if "learning checklist" in first:
            cur = "checklist"
        elif re.search(r'\brecap\b', first):
            cur = "recap"
        elif re.match(r'^(unit\s+)?reflect\b', first):
            cur = "reflect"
        buckets[cur].append(txt)
    row["Unit Learning Checklist"] = "\n".join(buckets["checklist"]).strip()
    row["Unit Quick Recap"] = "\n".join(buckets["recap"]).strip()
    row["Unit Reflect"] = "\n".join(buckets["reflect"]).strip()
    return row


# Portfolio docs are inconsistent across units: sections may be numbered
# "1)" or "1." or not numbered at all, and headings may sit on their own line
# (content following) or inline with their content. Sections are therefore
# matched by LABEL NAME (leading number marker stripped first). Each label
# regex consumes only the label; text after match.end() is inline content.
# Order matters -- more specific labels first.
_PORTFOLIO_SECTIONS = [
    (re.compile(r'^project name[:\-]?\s*', re.I),           "Project Title"),
    (re.compile(r'^purpose link[:\-]?\s*', re.I),           "Purpose"),
    (re.compile(r'^purpose[:\-]?\s*', re.I),                "Purpose"),
    (re.compile(r'^portfolio link[:\-]?\s*', re.I),         "Portfolio Link"),
    (re.compile(r'^what you will create[:\-]?\s*', re.I),   "What You Will Create"),
    (re.compile(r'^materials[:\-]?\s*', re.I),              "Materials"),
    (re.compile(r'^(?:numbered steps|steps?[- ]by[- ]steps?[^:\n]*|'
                r'steps to (?:create|complete)[^:\n]*|steps)[:\-]?\s*', re.I),
     "Steps"),
    (re.compile(r'^(?:templates?\s*(?:/|and)\s*organisers?[^:\n]*|'
                r'templates?\s+and\s+how to fill(?:\s+it)?|'
                r'templates?\b[^:\n]*?how to fill(?:\s+it)?|'
                r'how to fill(?:\s+it)?)[:\-]?\s*', re.I), "Template How To Fill"),
    (re.compile(r'^(?:completed example|example artefact|example)[:\-]?\s*', re.I),
     "Example"),
    (re.compile(r'^common mistakes?\s*(?:\+|and|&)?\s*fixes[:\-]?\s*', re.I),
     "Common Mistakes And Fixes"),
    (re.compile(r'^submission checklist[:\-]?\s*', re.I),   "Submission Checklist"),
    (re.compile(r'^optional extension[:\-]?\s*', re.I),     "Optional Extension"),
    (re.compile(r'^end-?of-?(?:unit|term)[^:\n]*inventory[:\-]?\s*', re.I),
     "End-of-Unit Inventory"),
    (re.compile(r'^motivation\s*(?:\+|and|&)?\s*reflection'
                r'(?:\s+skill badge earned)?[:\-]?\s*', re.I), "Motivation Reflection"),
    (re.compile(r'^[^\w]*skill badge earned[:\-]?\s*', re.I), "Motivation Reflection"),
    (re.compile(r'^reflect[^:\n]*[:\-]?\s*', re.I),         "Unit Reflect"),
]

# recognised but non-column headings: capturing stops so their body does not
# bleed into the previous column
_PORTFOLIO_STOP = re.compile(
    r'^(?:new tool\b|portfolio check\b|quick check rubric\b|'
    r'check:? my learning rubric\b|colour[- ]coding guide\b|the noun data box\b)',
    re.I)

_NUMPREFIX = re.compile(r'^\d{1,2}[.)]\s*')
# strip a "Skill Badge Earned:" lead plus any emoji/symbol that follows it
_RE_BADGE = re.compile(r'^[^\w]*skill badge earned\s*:?\s*[^\w\s]*\s*', re.I)


def _portfolio_lines(paras):
    """Flatten paragraphs to clean single lines. Headings and their content are
    sometimes in one paragraph split by a newline, sometimes in separate
    paragraphs -- splitting on newlines normalises both."""
    lines = []
    for p in paras:
        for ln in p.text.split("\n"):
            s = strip_bullet(ln.strip())
            if s:
                lines.append(s)
    return lines


def _match_portfolio_section(line):
    """Return (column, inline_content) if the line is a portfolio section
    heading, else None. A leading '1)'/'1.' number marker is ignored."""
    s = _NUMPREFIX.sub("", line)
    for rx, col in _PORTFOLIO_SECTIONS:
        m = rx.match(s)
        if m:
            return col, s[m.end():].strip()
    return None


def parse_portfolio(doc, path):
    """End of Unit N Portfolio -> the Overall Portfolio Project columns.
    Sections are matched by label name (any numbering style). Content on the
    lines following a heading wins; inline text after the label is the
    fallback for docs that keep label and content on one line."""
    paras = _nonempty_paras(doc)
    row = {"_file": os.path.basename(path)}
    row["Unit Number"] = _unit_from_name(path, [_RE_EOU_PORTFOLIO])

    out = {}
    cur_col, inline, following = None, "", []

    def flush():
        if cur_col is not None:
            parts = ([inline] if inline else []) + following
            content = "\n".join(parts).strip()
            if cur_col == "Motivation Reflection":
                content = _RE_BADGE.sub("", content).strip()
            # keep the first non-empty capture for a column; never blank it out
            if content and not out.get(cur_col):
                out[cur_col] = content

    for s in _portfolio_lines(paras[1:]):    # paras[0] is the doc title
        hit = _match_portfolio_section(s)
        if hit:
            flush()
            cur_col, inline, following = hit[0], hit[1], []
        elif _PORTFOLIO_STOP.match(_NUMPREFIX.sub("", s)):
            flush()
            cur_col, inline, following = None, "", []
        else:
            if cur_col is not None:
                following.append(s)
    flush()

    # Reflect sometimes rides inline on the Motivation line
    # ("...Reflection prompt: ..." / "...Reflect: ...") -- split it out.
    mot = out.get("Motivation Reflection", "")
    if mot and not out.get("Unit Reflect"):
        mk = re.search(r'\b(?:reflection prompt|reflect)\s*[:\-]\s*', mot, re.I)
        if mk:
            out["Motivation Reflection"] = mot[:mk.start()].strip()
            out["Unit Reflect"] = mot[mk.end():].strip()

    # project name is occasionally only in the document title line
    if not out.get("Project Title") and paras:
        t = re.sub(r'^\s*portfolio project\s*\d*\s*[:—–-]?\s*', "",
                   paras[0].text.strip(), flags=re.I).strip()
        if t:
            out["Project Title"] = t

    row.update(out)
    return row


def parse_docx(path, unknown_labels=None):
    """Parse one .docx into a wide row dict. Routes unit-level documents to
    their dedicated parser; ordinary lessons use the bold-label parser.
    unknown_labels: optional set to collect non-whitelisted bold labels."""
    if unknown_labels is None:
        unknown_labels = set()
    doc = Document(path)
    doc_type = classify_doc(path)

    if doc_type != "lesson":
        row = {"unit_intro": parse_unit_intro,
               "end_of_unit": parse_end_of_unit,
               "end_of_unit_portfolio": parse_portfolio}[doc_type](doc, path)
        row["_doc_type"] = doc_type
        all_text = "\n".join(p.text for p in doc.paragraphs)
        row.update(detect_images(path, all_text))
        return row

    return parse_lesson(doc, path, unknown_labels, doc_type)


def parse_lesson(doc, path, unknown_labels, doc_type="lesson"):
    """Original verbatim bold-label parser for ordinary sub-unit lessons."""
    title, sections = split_sections(doc, unknown_labels)
    row = {"_file": os.path.basename(path), "_doc_type": doc_type}

    m = re.match(r'^\s*(\d+(?:\.\d+)*)\s+(.*)$', title)
    if m:
        row["Unit Number"] = m.group(1)
        row["Lesson Title"] = m.group(2).strip()
    else:
        # title line may sit deeper in the doc (e.g. bold "1.2 The ...")
        row["Unit Number"] = ""
        row["Lesson Title"] = title
        for p in doc.paragraphs[:15]:
            m2 = re.match(r'^\s*(\d+(?:\.\d+)*)\s+(\S.*)$', p.text.strip())
            if m2:
                row["Unit Number"] = m2.group(1)
                row["Lesson Title"] = m2.group(2).strip()
                break
        if not row["Unit Number"]:
            # last resort: filename like "1.2 What is ...docx"
            m3 = re.match(r'^\s*(\d+(?:\.\d+)*)\s+(.*?)\.docx$',
                          os.path.basename(path), re.I)
            if m3:
                row["Unit Number"] = m3.group(1)

    misc = []
    for col, lines in sections:
        if col == "Mini Practice":
            for k, v in fan_mini_practice(lines).items():
                row[k] = v
            continue
        clean = [(l[len("[title] "):] if l.startswith("[title] ") else l) for l in lines]
        text = "\n".join(l for l in clean if l != "").strip()
        if col == "_misc":
            misc.append(text)
        else:
            row[col] = (row.get(col, "") + "\n" + text).strip() if row.get(col) else text

    if misc:
        # a stray copy of the title line is not real unplaced content
        title_line = f"{row['Unit Number']} {row['Lesson Title']}".strip()
        misc = [x for x in misc if x and x.strip() != title_line]
        if misc:
            row["_misc"] = " | ".join(misc)

    # recover the comprehension passage title when it wasn't captured as a
    # section (number-prefixed "N.N Main Comprehension Text" or a plain
    # "Poem:/Drama Scene:" line before BEFORE READING)
    if (row.get("Before Reading") or row.get("Main Text")) and \
            not row.get("Comprehension Title"):
        t = find_comprehension_title(doc)
        if t:
            row["Comprehension Title"] = t

    all_text = "\n".join(p.text for p in doc.paragraphs)
    row.update(detect_images(path, all_text))
    return row


def remap_unknown(row, mapping):
    """Move '?Label' columns onto canonical columns per judge mapping.
    mapping: {raw_label: canonical_column or None}. None keeps the raw label
    as its own column (surfaced, not hidden)."""
    for key in list(row.keys()):
        if not key.startswith("?"):
            continue
        raw = key[1:]
        target = mapping.get(raw)
        text = row.pop(key)
        # the heading itself is content too (often it IS the value, e.g. a
        # passage title mapped to Comprehension Title) -- keep it with its body
        if target and target != "IGNORE":
            text = (raw + "\n" + text).strip() if text else raw
            row[target] = (row.get(target, "") + "\n" + text).strip() if row.get(target) else text
        else:
            # includes IGNORE: never discard content -- verbatim guarantee
            row[raw] = text
    return row
