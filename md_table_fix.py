"""
Fix markdown pipe tables that were pasted into .docx documents as plain
paragraph text, replacing each with a real Word table at the same
position. Never modifies originals -- writes fixed copies to out_dir.

Verbatim by construction: cell text is sliced from the paragraph text,
never rewritten.
"""
import os
import re

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from main import find_docx

_SEP_TOKEN = re.compile(r"^:?-{3,}:?$")


def _is_sep_row(cells):
    """True if every cell in a row looks like a markdown separator cell."""
    return bool(cells) and all(_SEP_TOKEN.match(c.strip()) for c in cells)


def _row_cells(text):
    """Split a single pipe-row line into cells (drop leading/trailing
    empty tokens from leading/trailing '|')."""
    parts = [p.strip() for p in text.strip().split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return parts


def _looks_like_pipe_row(text):
    t = text.strip()
    return t.count("|") >= 2


# ---------- form B: single-paragraph concatenated table ----------

def _detect_form_b_safe(text):
    """Wrapper: returns ('ok', before, header, rows, after) or
    ('ragged', header, n_cols) or None (not a form-B table)."""
    if text.count("|") < 8:
        return None
    tokens = text.split("|")
    stripped = [t.strip() for t in tokens]
    sep_positions = [i for i, t in enumerate(stripped) if t != "" and _SEP_TOKEN.match(t)]
    if len(sep_positions) < 2:
        return None
    # find the longest consecutive run of separator tokens
    best_start = best_end = None
    i = 0
    while i < len(stripped):
        if stripped[i] != "" and _SEP_TOKEN.match(stripped[i]):
            j = i
            while j < len(stripped) and stripped[j] != "" and _SEP_TOKEN.match(stripped[j]):
                j += 1
            if best_start is None or (j - i) > (best_end - best_start):
                best_start, best_end = i, j
            i = j
        else:
            i += 1
    if best_start is None or (best_end - best_start) < 2:
        return None
    sep_run_start, sep_run_end = best_start, best_end
    n_cols = sep_run_end - sep_run_start

    header_tokens = []
    k = sep_run_start - 1
    while k >= 0 and len(header_tokens) < n_cols:
        tok = stripped[k]
        if tok != "":
            header_tokens.append(tok)
        k -= 1
    header_tokens.reverse()
    if len(header_tokens) != n_cols:
        return None
    header_start_idx = k + 1

    before_tokens = tokens[:header_start_idx]
    before_text = "|".join(before_tokens).strip()
    before_text = before_text.strip("|").strip()

    data_tokens_raw = stripped[sep_run_end:]
    data_tokens = [t for t in data_tokens_raw if t != ""]

    if len(data_tokens) % n_cols != 0:
        return ("ragged", header_tokens, n_cols)

    rows = [data_tokens[i:i + n_cols] for i in range(0, len(data_tokens), n_cols)]
    after_text = ""
    return ("ok", before_text, header_tokens, rows, after_text)


# ---------- form A: multi-line paragraphs ----------

def _detect_form_a_block(paragraphs, start):
    """Given a list of paragraph objects and a start index, try to parse a
    multi-line MD table beginning at start. Returns (end_index_exclusive,
    header, rows) or None."""
    if start + 1 >= len(paragraphs):
        return None
    p0 = paragraphs[start].text
    p1 = paragraphs[start + 1].text
    if not _looks_like_pipe_row(p0) or not _looks_like_pipe_row(p1):
        return None
    header = _row_cells(p0)
    sep_cells = _row_cells(p1)
    if not header or not _is_sep_row(sep_cells):
        return None
    n_cols = len(sep_cells)
    if len(header) != n_cols:
        return None
    rows = []
    i = start + 2
    while i < len(paragraphs):
        t = paragraphs[i].text
        if not _looks_like_pipe_row(t):
            break
        cells = _row_cells(t)
        if len(cells) != n_cols:
            break
        rows.append(cells)
        i += 1
    return i, header, rows


# ---------- Word table construction ----------

def _set_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _build_table(doc, anchor_paragraph, header, rows):
    n_cols = len(header)
    n_rows = 1 + len(rows)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        _set_table_borders(tbl)

    for c, text in enumerate(header):
        cell = tbl.cell(0, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            tbl.cell(r, c).text = text

    anchor_paragraph._p.addprevious(tbl._tbl)
    return tbl


def _insert_plain_paragraph(doc, anchor_paragraph, text):
    from docx.text.paragraph import Paragraph
    p_el = OxmlElement("w:p")
    anchor_paragraph._p.addprevious(p_el)
    new_p = Paragraph(p_el, anchor_paragraph._parent)
    if text:
        new_p.add_run(text)
    return new_p


def _process_doc(path, out_dir):
    doc = docx.Document(path)
    body = doc.paragraphs
    n_converted = 0
    warnings = []

    i = 0
    while i < len(body):
        para = body[i]
        text = para.text

        if not text or "|" not in text:
            i += 1
            continue

        # try form A first (multi-line)
        formA = _detect_form_a_block(body, i)
        if formA:
            end, header, rows = formA
            if rows:
                anchor = body[end] if end < len(body) else None
                if anchor is not None:
                    _build_table(doc, anchor, header, rows)
                else:
                    # table is the very last content -- append at end
                    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
                    try:
                        tbl.style = "Table Grid"
                    except KeyError:
                        _set_table_borders(tbl)
                    for c, h in enumerate(header):
                        cell = tbl.cell(0, c)
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(h)
                        run.bold = True
                    for r, row in enumerate(rows, start=1):
                        for c, v in enumerate(row):
                            tbl.cell(r, c).text = v
                # remove source paragraphs start..end-1
                for k in range(i, end):
                    p = body[k]
                    p._p.getparent().remove(p._p)
                n_converted += 1
                body = doc.paragraphs
                # i stays; re-fetch body, continue scanning from same index
                continue
            else:
                i += 1
                continue

        # try form B (single paragraph)
        res = _detect_form_b_safe(text)
        if res is None:
            i += 1
            continue
        if res[0] == "ragged":
            warnings.append(f"! md-table skipped (ragged rows): {os.path.basename(path)}")
            i += 1
            continue
        _, before_text, header, rows, after_text = res
        if not rows:
            i += 1
            continue

        anchor = para
        if before_text:
            _insert_plain_paragraph(doc, anchor, before_text)
        _build_table(doc, anchor, header, rows)
        if after_text:
            _insert_plain_paragraph(doc, anchor, after_text)

        p = para
        p._p.getparent().remove(p._p)
        n_converted += 1
        body = doc.paragraphs
        i += 1  # element removed; next index in refreshed body is same position

    for w in warnings:
        print(f"  {w}")

    if n_converted > 0:
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, os.path.basename(path))
        doc.save(out_path)
        print(f"  fixed {n_converted} md table(s): {os.path.basename(path)} -> "
              f"{out_path}")
    return n_converted


def fix_md_tables(sources, out_dir):
    files = find_docx(sources)
    results = []
    for path in files:
        try:
            n = _process_doc(path, out_dir)
        except Exception as e:
            print(f"  [!] md-table fix failed for {os.path.basename(path)}: {e}")
            continue
        if n > 0:
            results.append((path, n))
    return results
