"""
parser_it.py

Deterministic IT lesson .docx parser (Form 1 IT layout).

IT lessons are one document per module, structured as numbered sections:

    1  Learning Objectives
    2  Introduction / Topic Overview
    3  Key Concepts & Explanations   (with 3.1 / 3.2 ... sub-headings)
    4  Practical Activity (Hands-On)
    5  Guided Practice (With Support)
    6  Independent Practice (No Support)
    7  Challenge / Higher-Order Thinking
    8  Knowledge Check (Quick Assessment)
    9  Common Mistakes & Tips
    10 Real-World Application
    11 Summary / Key Takeaways
    12 Answers & Solutions
    13 Suggested Interactive Moments
    14 Developer Notes

Numbers vary between documents; sections are matched by NAME. Headings are
either Word Heading styles or fully-bold paragraphs. Extraction is verbatim:
text is sliced, never regenerated. Unknown numbered headings keep their own
'?Label' column so nothing is dropped.

Author: AbhishekAEDan
"""
__author__ = "AbhishekAEDan"

import os
import re

from docx import Document

from parser_core import bold_fraction, strip_bullet, detect_images

# canonical IT section names, matched by normalized *contains* keys.
# Order matters -- more specific keys first.
IT_SECTIONS = [
    ("learning objective",       "Learning Objectives"),
    ("topic overview",           "Introduction / Topic Overview"),
    ("introduction",             "Introduction / Topic Overview"),
    ("key concept",              "Key Concepts"),
    ("practical activity",       "Practical Activity"),
    ("guided practice",          "Guided Practice"),
    ("independent practice",     "Independent Practice"),
    ("higher-order",             "Challenge / Higher-Order Thinking"),
    ("higher order",             "Challenge / Higher-Order Thinking"),
    ("knowledge check",          "Knowledge Check"),
    ("quick assessment",         "Knowledge Check"),
    ("common mistake",           "Common Mistakes & Tips"),
    ("real-world application",   "Real-World Application"),
    ("real world application",   "Real-World Application"),
    ("key takeaway",             "Summary / Key Takeaways"),
    ("summary",                  "Summary / Key Takeaways"),
    ("answers & solutions",      "Answers & Solutions"),
    ("answers and solutions",    "Answers & Solutions"),
    ("answer & solution",        "Answers & Solutions"),
    ("answers/solutions",        "Answers & Solutions"),
    ("interactive moment",       "Suggested Interactive Moments"),
    ("developer note",           "Developer Notes"),
]

IT_ORDER = ["_file", "_doc_type", "Module Number", "Module Title",
            "Module Descriptive Title", "Form", "Term", "Practical Title",
            "Learning Objectives", "Introduction / Topic Overview",
            "Key Concepts", "Practical Activity", "Guided Practice",
            "Independent Practice", "Challenge / Higher-Order Thinking",
            "Knowledge Check", "Common Mistakes & Tips",
            "Real-World Application", "Summary / Key Takeaways",
            "Answers & Solutions", "Suggested Interactive Moments",
            "Developer Notes",
            "has_image", "image_missing", "image_note", "_misc"]

# top-level numbered heading: "6 Guided Practice", "5. Practical Activity",
# "10) Common Mistakes" -- but NOT "4.1 Sub heading" (decimal = sub-level)
_TOP_NUM = re.compile(r'^(\d{1,2})(?![.)]?\d|\.\d)[.)]?\s+(.*)$')
_SUB_NUM = re.compile(r'^\d{1,2}\.\d{1,2}\b')

_RE_MODULE = re.compile(r'\bmodule\s*(\d+(?:\.\d+)?)', re.I)
_RE_FORM = re.compile(r'\bform\s*(\d+)', re.I)
_RE_TERM = re.compile(r'\bterm\s*(\d+)', re.I)


def _norm(s):
    return re.sub(r'\s+', ' ', s).strip().lower()


# ---------- list formatting (bullets / numbering / indent) ----------

def _roman(n):
    out = ""
    for v, s in ((10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i")):
        while n >= v:
            out += s
            n -= v
    return out


class ListFormatter:
    """Reconstructs the visible list markers Word renders from numbering.xml
    (p.text carries none of them). Bullets become '•', numbered items get
    their counter ('1.', 'a)', 'i.'), nesting becomes leading tabs."""

    def __init__(self, doc):
        self._fmt_cache = {}
        self._counters = {}       # numId -> {ilvl: count}
        try:
            self._numbering = doc.part.numbering_part.element
        except (AttributeError, KeyError, NotImplementedError):
            self._numbering = None

    def _num_fmt(self, num_id, ilvl):
        key = (num_id, ilvl)
        if key in self._fmt_cache:
            return self._fmt_cache[key]
        fmt = "bullet"
        el = self._numbering
        if el is not None:
            try:
                aid = el.xpath(f'w:num[@w:numId="{num_id}"]/'
                               f'w:abstractNumId/@w:val')
                if aid:
                    vals = el.xpath(
                        f'w:abstractNum[@w:abstractNumId="{aid[0]}"]/'
                        f'w:lvl[@w:ilvl="{ilvl}"]/w:numFmt/@w:val')
                    if vals:
                        fmt = vals[0]
            except Exception:
                pass
        self._fmt_cache[key] = fmt
        return fmt

    def marker(self, p):
        """Return (prefix, ilvl) for a paragraph: '' when it is not a list
        item, else the rendered marker like '•' or '2.'"""
        pPr = p._p.pPr
        if pPr is None or pPr.numPr is None:
            return "", 0
        try:
            num_id = pPr.numPr.numId.val
            ilvl = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else 0
        except AttributeError:
            return "", 0
        fmt = self._num_fmt(num_id, ilvl)
        if fmt in ("bullet", "none"):
            return "•", ilvl
        c = self._counters.setdefault(num_id, {})
        c[ilvl] = c.get(ilvl, 0) + 1
        for lv in list(c):
            if lv > ilvl:               # deeper levels restart
                del c[lv]
        n = c[ilvl]
        if fmt == "lowerLetter":
            return f"{chr(96 + ((n - 1) % 26) + 1)})", ilvl
        if fmt == "upperLetter":
            return f"{chr(64 + ((n - 1) % 26) + 1)}.", ilvl
        if fmt == "lowerRoman":
            return f"{_roman(n)}.", ilvl
        if fmt == "upperRoman":
            return f"{_roman(n).upper()}.", ilvl
        return f"{n}.", ilvl            # decimal and friends


def _para_images(p):
    """Return ['[IMAGE: alt text]'] lines for pictures embedded in the
    paragraph (alt text kept when the author set one)."""
    out = []
    try:
        drawings = p._p.xpath('.//w:drawing')
    except Exception:
        return out
    for d in drawings:
        descr = d.xpath('.//wp:docPr/@descr') or d.xpath('.//wp:docPr/@name')
        note = (descr[0].strip() if descr and descr[0].strip() else "")
        out.append(f"[IMAGE: {note}]" if note else "[IMAGE]")
    return out


def match_it_section(text):
    """Return canonical section name if the heading text names a known IT
    section, else None."""
    n = _norm(text)
    for key, name in IT_SECTIONS:
        if key in n:
            return name
    # bare "Challenge" heading -- exact only, so content sub-titles like
    # "Quick Break Challenge" stay inside their section
    if n.rstrip(" :") == "challenge":
        return "Challenge / Higher-Order Thinking"
    return None


def _is_heading(p, txt):
    """Heading = Word Heading 1/2 style, or a fully-bold short line."""
    style = (p.style.name or "").lower()
    if style.startswith("heading"):
        # any heading level counts -- documents use H1..H3 for top sections;
        # N.N-numbered sub-headings are filtered by the caller
        return True
    return bold_fraction(p) >= 0.9 and len(txt) <= 90


# lesson number: "Lesson 2.4", "Lessons 3.6-3.9"
_RE_LESSON = re.compile(r'\blessons?\s*(\d+(?:\.\d+)*(?:\s*[-–]\s*\d+(?:\.\d+)*)?)',
                        re.I)
# "1. 2.6  Types of Software ..." -- outline number then a lesson number
_RE_LEAD_LESSON = re.compile(r'^\d{1,2}[.)]\s*(\d+\.\d+(?:\s*[-–]\s*\d+\.\d+)?)\s+')
# "Module 3: Word Processing" / "Module 4 — Internet and Web 2.0 Tools | ..."
_RE_THEME = re.compile(r'^module\s*\d+\s*[:—–]\s*([^|]+?)\s*(?:\||$)', re.I)


def _clean_title(line):
    """Best-effort descriptive title from a combined title line."""
    t = re.sub(r'\(\s*module[^)]*\)\s*$', '', line, flags=re.I).strip()
    t = _RE_LEAD_LESSON.sub("", t)
    # "... Term 1: <title>" or "... Term 1 <title>"
    m = re.search(r'term\s*\d+\s*[:,]?\s*(.+)$', t, re.I)
    if m and m.group(1).strip():
        t = m.group(1).strip()
    else:
        # "Lesson 2.4: <title>"
        m = re.match(r'^\s*lessons?\s*[\d.\-–\s]+[:—–-]\s*(.+)$', t, re.I)
        if m:
            t = m.group(1).strip()
        else:
            # bare leading outline number "1 " / "1. "
            t = re.sub(r'^\d{1,2}[.)]?\s+', '', t)
    return t.rstrip(" :-–—")


def parse_title_block(lines, path, row):
    """Fill Module Number / Module Title / Module Descriptive Title /
    Form / Term from the lines that precede the first section heading."""
    text = "\n".join(lines)
    module = _RE_MODULE.search(text)
    lesson = _RE_LESSON.search(text)
    if not lesson:
        for ln in lines:
            m = _RE_LEAD_LESSON.match(ln)
            if m:
                lesson = m
                break
    form = _RE_FORM.search(text)
    term = _RE_TERM.search(text)
    row["Module Number"] = (lesson.group(1).replace(" ", "") if lesson
                            else (module.group(1) if module else ""))
    row["Form"] = form.group(1) if form else ""
    row["Term"] = term.group(1) if term else ""

    theme, descriptive, best = "", "", ""
    for ln in lines:
        if ln.startswith("["):          # image/diagram markers are not titles
            continue
        m = _RE_THEME.match(ln)
        if m and not theme:
            theme = m.group(1).strip()
            continue
        if re.match(r'^lessons?\b', ln, re.I) and not descriptive:
            descriptive = ln.strip()
        if not best:
            c = _clean_title(ln)
            # skip lines that are pure module/form/term bookkeeping
            left = re.sub(r'\b(module|form|term)\s*\d+\b', '', c, flags=re.I)
            left = re.sub(r'[\d\s|,.:—–\-()]+', '', left)
            if left:
                best = c
    # sample sheets use the module theme as "Module Title 1" when present
    row["Module Title"] = theme or best or \
        os.path.splitext(os.path.basename(path))[0]
    if descriptive:
        row["Module Descriptive Title"] = descriptive


def parse_it_docx(path, unknown_labels=None):
    """Parse one IT lesson .docx into a wide row dict. Sections keep their
    verbatim text joined with newlines. unknown_labels: optional set that
    collects numbered headings that match no known section."""
    if unknown_labels is None:
        unknown_labels = set()
    doc = Document(path)
    row = {"_file": os.path.basename(path), "_doc_type": "it_lesson"}

    def _has_drawing(p):
        try:
            return bool(p._p.xpath('.//w:drawing'))
        except Exception:
            return False

    paras = [p for p in doc.paragraphs if p.text.strip() or _has_drawing(p)]

    sections = {}          # canonical name -> [lines]
    title_block = []       # lines before the first section heading
    misc = []
    current = None

    def add(line):
        if current is None:
            title_block.append(line)
        else:
            sections.setdefault(current, []).append(line)

    listfmt = ListFormatter(doc)

    for p in paras:
        style = (p.style.name or "").lower()
        marker, ilvl = listfmt.marker(p)
        first_line = True
        for ln in p.text.strip().split("\n"):
            raw_line = ln.rstrip()
            txt = strip_bullet(ln.strip())
            if not txt:
                continue
            # visible list marker + nesting indent, on the first line only
            if marker and first_line:
                fmt_line = "\t" * ilvl + f"{marker} {txt}"
            else:
                fmt_line = raw_line if raw_line.strip() else txt
            first_line = False
            # authoring junk above the content: comments and cover-tab titles
            if current is None and (txt.startswith("//") or style == "title"):
                continue
            # a line naming the module/lesson before any section began is
            # part of the title block, never a section heading
            if current is None and (_RE_MODULE.search(txt)
                                    or _RE_LESSON.search(txt)):
                title_block.append(txt)
                continue
            hit = None
            m = _TOP_NUM.match(txt)
            # list items are never section headings
            heading_p = not marker and _is_heading(p, txt)
            # an unnumbered bold label ending with ':' inside a section is a
            # sub-label ("Guided Practice:" inside Answers & Solutions), not
            # a new section
            sub_label = (current is not None and not m
                         and txt.rstrip().endswith(":"))
            if heading_p and not _SUB_NUM.match(txt) and not sub_label:
                # numbered heading or a styled/bold heading naming a section
                cand = m.group(2) if m else txt
                hit = match_it_section(cand)
                # a numbered STYLE heading (Heading 1/2) that matches nothing
                # is an unknown section; plain-bold numbered lines inside a
                # section are content (e.g. "1. Dust" in Key Concepts)
                if (hit is None and m and len(cand) <= 70
                        and (style.startswith("heading") or current is None)):
                    unknown_labels.add(cand.strip())
                    current = "?" + cand.strip()
                    sections.setdefault(current, [])
                    continue
            # inside Answers & Solutions, plain-bold restatements of earlier
            # sections ("Guided Practice Answers", "5. Guided Practice:",
            # "Section 5: Practical Activity") are answer sub-blocks, not new
            # sections -- only real Word Heading styles may switch away
            if (hit and current == "Answers & Solutions"
                    and not style.startswith("heading")
                    and hit not in ("Suggested Interactive Moments",
                                    "Developer Notes")):
                sections[current].append(fmt_line)
                continue
            if hit:
                current = hit
                sections.setdefault(current, [])
                # heading may carry an inline title, e.g.
                # "5. Practical Activity (Hands-On): Correct Posture ..."
                rest = ""
                mm = re.search(r'\)\s*[:—–-]\s*(.+)$', txt)
                if mm:
                    rest = mm.group(1).strip()
                if hit == "Practical Activity" and rest:
                    row["Practical Title"] = rest
                elif rest:
                    sections[current].append(rest)
                continue
            add(fmt_line)
        # pictures embedded in this paragraph stay with its section
        for img in _para_images(p):
            add(img)

    parse_title_block(title_block, path, row)

    for name, lines in sections.items():
        text = "\n".join(lines).strip()
        if text or not name.startswith("?"):
            row[name] = text
    if misc:
        row["_misc"] = " | ".join(m for m in misc if m)

    all_text = "\n".join(p.text for p in doc.paragraphs)
    row.update(detect_images(path, all_text))
    return row


# ---------- subject auto-detection ----------

_ELA_MARKERS = ["vocab vault", "mini practice", "my goal", "mistake spotter",
                "bridge back", "check for understanding", "quick recap"]
_IT_MARKERS = ["key concepts", "knowledge check", "guided practice",
               "independent practice", "practical activity",
               "higher-order thinking", "real-world application",
               "answers & solutions", "topic overview"]


def detect_subject_for_file(path):
    """Return 'ela', 'it', or '' for one document, by counting the section
    markers each layout uses."""
    try:
        doc = Document(path)
    except Exception:
        return ""
    text = "\n".join(p.text for p in doc.paragraphs[:120]).lower()
    ela = sum(1 for k in _ELA_MARKERS if k in text)
    it = sum(1 for k in _IT_MARKERS if k in text)
    if it > ela:
        return "it"
    if ela > it:
        return "ela"
    return ""


def detect_subject(files, sample=4):
    """Detect the subject across up to `sample` documents.
    Returns ('ela'|'it'|'', votes_dict)."""
    votes = {"ela": 0, "it": 0}
    for f in files[:sample]:
        v = detect_subject_for_file(f)
        if v:
            votes[v] += 1
    if votes["it"] > votes["ela"]:
        return "it", votes
    if votes["ela"] > votes["it"]:
        return "ela", votes
    return "", votes
